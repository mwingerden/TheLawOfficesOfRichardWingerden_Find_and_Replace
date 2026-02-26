from pathlib import Path

import pytest
from docx import Document

from find_and_replace import FindAndReplace


# ----------------------------
# Helpers
# ----------------------------

def create_docx(path: Path, text: str):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


# ----------------------------
# Tests
# ----------------------------

def test_folder_not_found():
    service = FindAndReplace()

    report = service.find_and_replace(
        source_folder="does_not_exist",
        word_find="A",
        word_replace="B",
    )

    assert report.success is False
    assert "Folder not found" in report.error_message


def test_no_docx_files(tmp_path):
    service = FindAndReplace()

    report = service.find_and_replace(
        source_folder=str(tmp_path),
        word_find="A",
        word_replace="B",
    )

    assert report.success is False
    assert report.error_message == "No .docx files found."


def test_blank_word_find_raises(tmp_path):
    service = FindAndReplace()

    with pytest.raises(ValueError):
        service.find_and_replace(
            source_folder=str(tmp_path),
            word_find="",
            word_replace="B",
        )


def test_blank_word_replace_raises(tmp_path):
    service = FindAndReplace()

    with pytest.raises(ValueError):
        service.find_and_replace(
            source_folder=str(tmp_path),
            word_find="A",
            word_replace="",
        )


def test_successful_replacement(tmp_path):
    # Must match one of the allowed prefixes in _file_order
    file_path = tmp_path / "Trust Summary.docx"

    create_docx(file_path, "Hello TEST world TEST")

    service = FindAndReplace()

    report = service.find_and_replace(
        source_folder=str(tmp_path),
        word_find="TEST",
        word_replace="DONE",
    )

    assert report.success is True
    assert report.files_processed == 1
    assert report.files_modified == 1
    assert report.total_replacements == 2

    # Verify content was updated
    updated_doc = Document(str(file_path))
    full_text = "\n".join(p.text for p in updated_doc.paragraphs)

    assert "TEST" not in full_text
    assert full_text.count("DONE") == 2


def test_file_not_modified_if_no_match(tmp_path):
    file_path = tmp_path / "Trust Summary.docx"
    create_docx(file_path, "Hello world")

    service = FindAndReplace()

    report = service.find_and_replace(
        source_folder=str(tmp_path),
        word_find="NOTFOUND",
        word_replace="DONE",
    )

    assert report.success is True
    assert report.files_processed == 1
    assert report.files_modified == 0
    assert report.total_replacements == 0


def test_ignores_non_matching_prefix(tmp_path):
    # This file does NOT start with a valid prefix
    file_path = tmp_path / "Random File.docx"
    create_docx(file_path, "TEST")

    service = FindAndReplace()

    report = service.find_and_replace(
        source_folder=str(tmp_path),
        word_find="TEST",
        word_replace="DONE",
    )

    assert report.success is False
    assert report.error_message == "No .docx files found."
